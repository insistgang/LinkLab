from __future__ import annotations

import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont

try:
    import cairosvg
except Exception:  # pragma: no cover - optional local dependency
    cairosvg = None

try:
    from opencc import OpenCC
except Exception:  # pragma: no cover - optional local dependency
    OpenCC = None

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
DOC_DIR = ROOT / "docs" / "competition_ai_x"
ASSET_DIR = DOC_DIR / "assets"
WORD_DIR = DOC_DIR / "word"
BRAND_DIR = ROOT / "linklab" / "assets" / "brand"

PURPLE = (118, 44, 230)
PURPLE_DARK = (72, 31, 147)
LIME = (199, 254, 83)
MINT = (66, 230, 178)
CYAN = (40, 205, 210)
INK = (20, 34, 36)
MUTED = (87, 104, 98)
PAPER = (248, 255, 241)
WHITE = (255, 255, 255)
WARNING = (233, 58, 81)
GOLD = (255, 179, 48)


def font_path(name: str) -> str:
    candidates = [
        Path(r"C:\Windows\Fonts") / name,
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for item in candidates:
        if item.exists():
            return str(item)
    return "arial.ttf"


FONT_REG = font_path("msyh.ttc")
FONT_BOLD = font_path("msyhbd.ttc")


def f(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size)


def gradient(size: tuple[int, int], start: tuple[int, int, int], end: tuple[int, int, int]) -> Image.Image:
    w, h = size
    img = Image.new("RGB", size, start)
    px = img.load()
    for y in range(h):
        t = y / max(h - 1, 1)
        for x in range(w):
            t2 = (t * 0.65) + (x / max(w - 1, 1) * 0.35)
            color = tuple(int(start[i] * (1 - t2) + end[i] * t2) for i in range(3))
            px[x, y] = color
    return img


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.splitlines():
        if not raw:
            lines.append("")
            continue
        current = ""
        tokens = re.findall(r"[A-Za-z0-9_./+<>=:%-]+|\s+|.", raw)
        for token in tokens:
            if token.isspace() and not current:
                continue
            candidate = current + token
            if text_size(draw, candidate, font)[0] <= max_width or not current:
                current = candidate
            else:
                lines.append(current.rstrip())
                current = token.lstrip()
                if text_size(draw, current, font)[0] > max_width:
                    carry = ""
                    for ch in current:
                        candidate = carry + ch
                        if text_size(draw, candidate, font)[0] <= max_width or not carry:
                            carry = candidate
                        else:
                            lines.append(carry)
                            carry = ch
                    current = carry
        if current:
            lines.append(current.rstrip())
    return lines


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: tuple[int, int, int],
    max_width: int,
    line_gap: int = 10,
) -> int:
    x, y = xy
    for line in wrap_text(draw, text, font, max_width):
        draw.text((x, y), line, font=font, fill=fill)
        y += text_size(draw, line or "字", font)[1] + line_gap
    return y


def rounded(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], radius: int, fill, outline=None, width: int = 1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill=PURPLE_DARK, width: int = 6):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    pts = [
        end,
        (int(end[0] - size * math.cos(angle - math.pi / 7)), int(end[1] - size * math.sin(angle - math.pi / 7))),
        (int(end[0] - size * math.cos(angle + math.pi / 7)), int(end[1] - size * math.sin(angle + math.pi / 7))),
    ]
    draw.polygon(pts, fill=fill)


def render_logo(size: int = 170) -> Image.Image:
    svg = BRAND_DIR / "logo.svg"
    out = ASSET_DIR / "logo_rendered.png"
    if cairosvg and svg.exists():
        cairosvg.svg2png(url=str(svg), write_to=str(out), output_width=size, output_height=size)
        return Image.open(out).convert("RGBA")
    img = Image.new("RGBA", (size, size), (199, 254, 83, 255))
    d = ImageDraw.Draw(img)
    rounded(d, (12, 12, size - 12, size - 12), 30, PURPLE)
    d.text((size * 0.24, size * 0.34), "Link", font=f(int(size * 0.18), True), fill=WHITE)
    return img


def draw_pill(draw, xy, label, fill, text_fill=WHITE, icon=None, font_size=30):
    x, y = xy
    font = f(font_size, True)
    tw, th = text_size(draw, label, font)
    w = tw + 74
    h = th + 30
    rounded(draw, (x, y, x + w, y + h), h // 2, fill)
    if icon:
        draw.text((x + 24, y + 13), icon, font=f(font_size - 4, True), fill=text_fill)
        tx = x + 52
    else:
        tx = x + 35
    draw.text((tx, y + 12), label, font=font, fill=text_fill)


def make_cover():
    w, h = 1800, 1050
    img = gradient((w, h), LIME, CYAN).convert("RGBA")
    d = ImageDraw.Draw(img)

    d.ellipse((-180, -260, 760, 680), fill=(242, 255, 205, 120))
    d.ellipse((1120, 610, 2040, 1380), fill=(70, 224, 170, 110))
    d.ellipse((1320, -260, 2180, 560), fill=(235, 255, 205, 110))

    logo = render_logo(150)
    img.alpha_composite(logo, (120, 100))
    d.text((300, 105), "共感 LinkAble", font=f(72, True), fill=INK)
    d.text((302, 200), "面向无障碍社会互助的 AI+X 服务智能体", font=f(36, True), fill=PURPLE_DARK)
    d.text((304, 268), "AI 第一响应 / 真人志愿者兜底 / SOS 安全闭环", font=f(28), fill=(42, 74, 68))

    draw_pill(d, (1220, 112), "赛题三 AI+X", PURPLE)
    draw_pill(d, (1220, 190), "竞赛 MVP", WHITE, PURPLE_DARK)

    card = (140, 410, 1660, 910)
    rounded(d, card, 44, (255, 255, 255, 224), outline=(232, 250, 228, 255), width=3)
    d.text((220, 455), "从“问答工具”", font=f(43, True), fill=INK)
    d.text((220, 515), "到“公共服务调度智能体”", font=f(43, True), fill=INK)
    body = "用户只需表达需求，AI 先完成意图识别、风险判断和标准化处理；当 AI 不确定、需求复杂或出现紧急信号时，系统自动转接真人志愿者或进入 SOS 流程。"
    draw_wrapped(d, (220, 605), body, f(30), (39, 76, 68), 670, 13)

    cx, cy = 1230, 650
    d.ellipse((cx - 110, cy - 110, cx + 110, cy + 110), fill=PURPLE)
    d.text((cx - 68, cy - 30), "AI\nAgent", font=f(34, True), fill=WHITE, align="center")
    nodes = [
        ("文字/语音/拍照", 1010, 470, MINT),
        ("直接回答", 1450, 470, LIME),
        ("志愿者匹配", 1450, 780, GOLD),
        ("SOS", 1010, 780, WARNING),
    ]
    for label, x, y, color in nodes:
        rounded(d, (x - 110, y - 44, x + 110, y + 44), 26, color)
        tw, th = text_size(d, label, f(25, True))
        d.text((x - tw / 2, y - th / 2 - 2), label, font=f(25, True), fill=INK if color != WARNING else WHITE)
        arrow(d, (x + (70 if x < cx else -70), y), (cx - (115 if x < cx else -115), cy + (0 if y == cy else (45 if y < cy else -45))), fill=(87, 73, 170), width=4)

    d.text((140, 965), "材料已完成匿名化检查，不含身份线索", font=f(24), fill=(38, 71, 66))
    img.convert("RGB").save(ASSET_DIR / "cover_banner.png", quality=95)


def make_architecture():
    w, h = 1800, 1200
    img = Image.new("RGB", (w, h), (251, 255, 247))
    d = ImageDraw.Draw(img)
    d.text((90, 70), "系统架构：AI 第一响应 + 真人互助 + SOS 闭环", font=f(48, True), fill=INK)
    d.text((92, 135), "UI 不直接调用具体 AI/OCR/WebRTC 服务，统一通过 facade 进入 Demo/Mock/Real 适配层。", font=f(28), fill=MUTED)

    input_nodes = [
        ("文字", "问题描述/说明书"),
        ("语音", "语音转文字"),
        ("拍照", "OCR/场景识别"),
        ("位置", "匹配/SOS 授权"),
    ]
    x0 = 120
    for i, (title, sub) in enumerate(input_nodes):
        x = x0 + i * 405
        rounded(d, (x, 230, x + 310, 390), 28, WHITE, outline=(204, 233, 216), width=3)
        d.text((x + 32, 260), title, font=f(34, True), fill=PURPLE_DARK)
        d.text((x + 32, 318), sub, font=f(23), fill=MUTED)
        arrow(d, (x + 155, 390), (900, 500), fill=(118, 44, 230), width=4)

    rounded(d, (520, 500, 1280, 750), 42, (238, 230, 255), outline=PURPLE, width=4)
    d.text((610, 545), "AgentServiceFacade", font=f(44, True), fill=PURPLE_DARK)
    for i, item in enumerate(["意图识别", "紧急度判断", "置信度评估", "安全文案", "下一步动作"]):
        draw_pill(d, (610 + (i % 3) * 210, 625 + (i // 3) * 66), item, PURPLE, font_size=24)

    outputs = [
        ("AI 直接解决", "OCR、药品说明、场景描述、颜色识别", LIME),
        ("志愿者匹配", "技能、距离、信誉、历史信任排序", GOLD),
        ("SOS 闭环", "撤销、广播、联系人通知、完成态", WARNING),
    ]
    for i, (title, sub, color) in enumerate(outputs):
        x = 165 + i * 540
        y = 900
        arrow(d, (900, 750), (x + 220, y - 30), fill=PURPLE_DARK, width=5)
        rounded(d, (x, y, x + 440, y + 170), 32, color)
        d.text((x + 35, y + 34), title, font=f(34, True), fill=WHITE if color == WARNING else INK)
        draw_wrapped(d, (x + 35, y + 92), sub, f(24), WHITE if color == WARNING else (45, 61, 54), 360, 8)

    img.save(ASSET_DIR / "system_architecture.png", quality=95)


def make_decision_flow():
    w, h = 1800, 1200
    img = Image.new("RGB", (w, h), (250, 255, 246))
    d = ImageDraw.Draw(img)
    d.text((90, 70), "AI 决策流程：不是强行回答，而是决定安全下一步", font=f(48, True), fill=INK)
    d.text((92, 135), "同一输入经过意图、紧急度、置信度和安全规则后，进入回答、追问、转人工或 SOS。", font=f(28), fill=MUTED)

    steps = [
        ("用户输入", "文字 / 语音 / 图片 / 位置"),
        ("识别意图", "OCR、场景、药品、导航、紧急等"),
        ("判断风险", "普通 / 较急 / 紧急"),
        ("评估置信度", "低于阈值不直接给确定答案"),
        ("选择动作", "answer / ask / match / SOS"),
    ]
    for i, (title, sub) in enumerate(steps):
        x = 115 + i * 335
        rounded(d, (x, 260, x + 270, 420), 28, WHITE, outline=(217, 236, 218), width=3)
        d.text((x + 35, 292), title, font=f(31, True), fill=PURPLE_DARK)
        draw_wrapped(d, (x + 35, 348), sub, f(20), MUTED, 205, 7)
        if i < len(steps) - 1:
            arrow(d, (x + 270, 340), (x + 325, 340), fill=PURPLE, width=5)

    lanes = [
        ("AI 可处理", "直接回答 + 朗读 + 继续追问", LIME),
        ("信息不足", "追问澄清，不装作确定", MINT),
        ("需要真人", "进入匹配，推荐技能标签", GOLD),
        ("紧急风险", "SOS 撤销、广播、联系人通知", WARNING),
    ]
    for i, (title, sub, color) in enumerate(lanes):
        x = 135 + i * 415
        y = 675
        rounded(d, (x, y, x + 330, y + 250), 36, color)
        d.text((x + 40, y + 45), title, font=f(33, True), fill=WHITE if color == WARNING else INK)
        draw_wrapped(d, (x + 40, y + 115), sub, f(25), WHITE if color == WARNING else (42, 64, 58), 245, 9)
        arrow(d, (900, 420), (x + 165, y), fill=(93, 69, 179), width=4)

    d.text((135, 1020), "关键边界：confidence < 0.65 不给确定答案；urgency = emergency 必须触发 SOS；can_resolve_by_ai = false 必须给出 handoff_reason。", font=f(26), fill=PURPLE_DARK)
    img.save(ASSET_DIR / "ai_decision_flow.png", quality=95)


def mini_phone(draw, box, title, lines, accent):
    x1, y1, x2, y2 = box
    rounded(draw, box, 36, (22, 33, 36), width=2)
    rounded(draw, (x1 + 12, y1 + 12, x2 - 12, y2 - 12), 28, (235, 255, 224))
    draw.rectangle((x1 + 12, y1 + 12, x2 - 12, y1 + 90), fill=accent)
    draw.text((x1 + 36, y1 + 35), title, font=f(24, True), fill=WHITE if accent != LIME else INK)
    yy = y1 + 130
    for line in lines:
        rounded(draw, (x1 + 35, yy, x2 - 35, yy + 78), 22, WHITE)
        draw_wrapped(draw, (x1 + 55, yy + 20), line, f(20, True), INK, x2 - x1 - 115, 4)
        yy += 98


def make_storyboard():
    w, h = 1800, 1050
    img = gradient((w, h), (238, 255, 218), (224, 253, 255))
    d = ImageDraw.Draw(img)
    d.text((90, 70), "5 分钟视频故事板", font=f(52, True), fill=INK)
    d.text((92, 140), "按“AI 处理 -> 转人工 -> 语音协助 -> SOS -> 指标”展示，而不是堆页面。", font=f(28), fill=MUTED)

    phones = [
        ("首页", ["AI 先处理", "复杂需求转真人"], PURPLE),
        ("AI 助手", ["布洛芬是什么药？", "可朗读 / 可转人工"], (83, 194, 132)),
        ("匹配", ["推荐志愿者", "技能：药品说明"], GOLD),
        ("通话", ["已接通", "语音协助中"], CYAN),
        ("SOS", ["10 秒撤销", "模拟通知联系人"], WARNING),
    ]
    start_x = 80
    for i, (title, lines, accent) in enumerate(phones):
        x = start_x + i * 345
        mini_phone(d, (x, 240, x + 285, 860), title, lines, accent)
        if i < len(phones) - 1:
            arrow(d, (x + 290, 545), (x + 330, 545), fill=PURPLE_DARK, width=5)

    draw_pill(d, (118, 910), "00:00-00:25 讲清定位", PURPLE, font_size=25)
    draw_pill(d, (500, 910), "00:25-02:00 AI 第一响应", PURPLE, font_size=25)
    draw_pill(d, (930, 910), "02:00-03:40 真人与 SOS 闭环", PURPLE, font_size=25)
    draw_pill(d, (1390, 910), "03:40-04:55 架构与指标", PURPLE, font_size=25)
    img.save(ASSET_DIR / "demo_storyboard.png", quality=95)


def make_metrics():
    w, h = 1800, 1050
    img = Image.new("RGB", (w, h), (249, 255, 245))
    d = ImageDraw.Draw(img)
    d.text((90, 70), "AI Inference 与工程验证指标", font=f(52, True), fill=INK)
    d.text((92, 140), "评测重点不是“回答是否像样”，而是意图、风险、动作和无障碍闭环是否可靠。", font=f(28), fill=MUTED)

    cards = [
        ("意图识别", ">=85%", "9 类主意图"),
        ("紧急召回", ">=95%", "SOS 高风险样本"),
        ("紧急误报", "<=2%", "普通需求对照"),
        ("AI 响应", "<=3s", "答案或下一步动作"),
        ("匹配耗时", "<=500ms", "50 人 Top 5"),
        ("无障碍", "200%", "字体缩放抽检"),
    ]
    for i, (title, value, sub) in enumerate(cards):
        x = 110 + (i % 3) * 560
        y = 250 + (i // 3) * 250
        rounded(d, (x, y, x + 500, y + 190), 34, WHITE, outline=(218, 238, 218), width=3)
        d.text((x + 36, y + 32), title, font=f(31, True), fill=PURPLE_DARK)
        d.text((x + 36, y + 82), value, font=f(48, True), fill=INK)
        d.text((x + 36, y + 148), sub, font=f(23), fill=MUTED)

    chart_x, chart_y = 170, 790
    d.text((chart_x, chart_y - 58), "提交前实测记录建议：", font=f(30, True), fill=INK)
    bars = [("intent", 85, PURPLE), ("SOS recall", 95, WARNING), ("action", 85, CYAN), ("accessibility", 100, MINT)]
    for i, (name, val, color) in enumerate(bars):
        y = chart_y + i * 48
        d.text((chart_x, y), name, font=f(22, True), fill=MUTED)
        rounded(d, (chart_x + 190, y + 4, chart_x + 690, y + 30), 13, (230, 239, 231))
        rounded(d, (chart_x + 190, y + 4, chart_x + 190 + int(val * 5), y + 30), 13, color)
        d.text((chart_x + 720, y), f"{val}%", font=f(22, True), fill=INK)

    rounded(d, (1030, 780, 1665, 980), 34, (235, 229, 255), outline=PURPLE, width=3)
    d.text((1070, 820), "证据包建议", font=f(30, True), fill=PURPLE_DARK)
    draw_wrapped(d, (1070, 875), "flutter analyze / flutter test 日志、APK hash、90 条样本评测表、3 分钟录屏、读屏与 200% 字体截图。", f(24), INK, 540, 10)
    img.save(ASSET_DIR / "metrics_dashboard.png", quality=95)


def generate_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    make_cover()
    make_architecture()
    make_decision_flow()
    make_storyboard()
    make_metrics()


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run, 9.5 if not bold else 10)


def set_run_font(run, size: float | None = None, color: RGBColor | None = None):
    run.font.name = "Microsoft YaHei"
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for name, size, color in [
        ("Heading 1", 18, PURPLE_DARK),
        ("Heading 2", 15, PURPLE_DARK),
        ("Heading 3", 12.5, INK),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return doc


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            text = row[c] if c < len(row) else ""
            set_cell_text(cell, text, bold=(r == 0))
            if r == 0:
                set_cell_shading(cell, "EADFFF")
            elif r % 2 == 0:
                set_cell_shading(cell, "F8FFF2")
    doc.add_paragraph()


def add_code_block(doc: Document, code: str):
    for raw in code.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(raw)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        set_cell_shading_paragraph(p, "F3F4F6")


def set_cell_shading_paragraph(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_markdown_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            set_run_font(run, 9.5)
        else:
            run = p.add_run(part)
            set_run_font(run, 10.5)
    return p


def md_to_docx(md_path: Path, docx_path: Path):
    converter = OpenCC("t2s") if OpenCC else None
    text = md_path.read_text(encoding="utf-8")
    if converter:
        text = converter.convert(text)
    lines = text.splitlines()
    doc = setup_doc()

    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, parse_table(block))
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            alt, rel = image_match.groups()
            img_path = (md_path.parent / rel).resolve()
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(6.5))
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                set_run_font(cap.runs[0], 9, RGBColor(87, 104, 98))
            i += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            doc.add_heading(heading.group(2), level=level)
            i += 1
            continue

        if stripped.startswith("- [ ]"):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(stripped[5:].strip())
            set_run_font(run, 10.5)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(stripped[2:].strip())
            set_run_font(run, 10.5)
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(numbered.group(1))
            set_run_font(run, 10.5)
            i += 1
            continue

        if stripped.startswith(">"):
            p = add_markdown_paragraph(doc, stripped.lstrip("> ").strip())
            p.paragraph_format.left_indent = Cm(0.5)
            set_cell_shading_paragraph(p, "F3F8EA")
            i += 1
            continue

        add_markdown_paragraph(doc, stripped)
        i += 1

    doc.save(docx_path)


def generate_word():
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    for md_path in sorted(DOC_DIR.glob("*.md")):
        docx_path = WORD_DIR / f"{md_path.stem}.docx"
        md_to_docx(md_path, docx_path)
        print(f"wrote {docx_path.relative_to(ROOT)}")


def main():
    generate_assets()
    generate_word()


if __name__ == "__main__":
    main()
